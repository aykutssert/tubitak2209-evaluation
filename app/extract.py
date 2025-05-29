from docx import Document

def extract_full_2209a_structure(file_path):
    doc = Document(file_path)
    result = {}
    current_main_section = None
    current_sub_section = None

    body_elements = list(doc.element.body.iterchildren())
    tables = iter(doc.tables)
    para_index = 0

    for elem in body_elements:
        if elem.tag.endswith('p'):
            para = doc.paragraphs[para_index]
            para_index += 1
            text = para.text.strip()
            text_upper = text.upper()

            # ALT BAŞLIKLAR - ÖZGÜN DEĞER ALTINDA
            if current_main_section == "ÖZGÜN DEĞER":
                if "1.1" in text or "KONUNUN ÖNEMİ" in text_upper:
                    current_sub_section = "1.1 Konunun Önemi, Özgün Değeri ve Hipotezi"
                    continue
                elif "1.2" in text or "AMAÇ" in text_upper:
                    current_sub_section = "1.2 Amaç ve Hedefler"
                    continue

            # ANA BAŞLIKLAR
            if "GENEL BİLGİLER" in text_upper:
                current_main_section = "GENEL BİLGİLER"
                current_sub_section = None
                result[current_main_section] = {}

            elif "ÖZET" in text_upper:
                current_main_section = "ÖZET"
                current_sub_section = None
                result[current_main_section] = {}

            elif "ÖZGÜN DEĞER" in text_upper and current_main_section != "ÖZGÜN DEĞER":
                current_main_section = "ÖZGÜN DEĞER"
                current_sub_section = None
                result[current_main_section] = {}

            elif "YÖNTEM" in text_upper and current_main_section != "YÖNTEM":
                current_main_section = "YÖNTEM"
                current_sub_section = None
                result[current_main_section] = {}
            elif "PROJE YÖNETİMİ" in text_upper and current_main_section != "PROJE YÖNETİMİ":
                current_main_section = "PROJE YÖNETİMİ"
                current_sub_section = None
                result[current_main_section] = {}
            elif "YAYGIN ETKİ" in text_upper and current_main_section != "4. YAYGIN ETKİ":
                current_main_section = "4. YAYGIN ETKİ"
                current_sub_section = None
                result[current_main_section] = {}

            elif "BÜTÇE TALEP ÇİZELGESİ" in text_upper:
                current_main_section = "BÜTÇE TALEP ÇİZELGESİ"
                current_sub_section = None
                result[current_main_section] = {}



        elif elem.tag.endswith('tbl') and current_main_section:
            table = next(tables)

            if current_main_section == "GENEL BİLGİLER":
                for row in table.rows:
                    if len(row.cells) == 1 and ":" in row.cells[0].text:
                        key, val = row.cells[0].text.split(":", 1)
                        result[current_main_section][key.strip()] = val.strip()

            elif current_main_section == "ÖZET":
                for row in table.rows:
                    if len(row.cells) == 1:
                        raw = row.cells[0].text.strip()
                        if "Anahtar Kelimeler" in raw:
                            result[current_main_section]["Anahtar Kelimeler"] = raw.split(":", 1)[-1].strip()
                        else:
                            result[current_main_section]["Özet"] = raw

            elif current_main_section == "ÖZGÜN DEĞER" and current_sub_section:
                content = ""
                for row in table.rows:
                    for cell in row.cells:
                        raw = cell.text.strip()
                        content += raw + "\n"
                result[current_main_section][current_sub_section] = content.strip()
                current_sub_section = None

            elif current_main_section == "YÖNTEM":
                if "Yöntem Açıklaması" in result[current_main_section]:
                    continue  # sadece ilk tablo alınır
                content = ""
                for row in table.rows:
                    for cell in row.cells:
                        raw = cell.text.strip()
                        content += raw + "\n"
                result[current_main_section]["Yöntem Açıklaması"] = content.strip()

            elif current_main_section == "PROJE YÖNETİMİ":
                # 3.1 İş-Zaman Çizelgesi
                if "3.1 İş-Zaman Çizelgesi" not in result[current_main_section]:
                    headers = []
                    rows_data = []

                    for i, row in enumerate(table.rows):
                        cells = [
                            ''.join(t.text for t in cell._element.iter() if t.tag.endswith('}t')).strip()
                            for cell in row.cells
                        ]
                        if i == 0:
                            headers = cells
                        else:
                            row_dict = {}
                            for h, v in zip(headers, cells):
                                row_dict[h] = v
                            rows_data.append(row_dict)

                    result[current_main_section]["3.1 İş-Zaman Çizelgesi"] = rows_data

                # 3.2 Risk Yönetimi
                elif "3.2 Risk Yönetimi" not in result[current_main_section]:
                    headers = []
                    rows_data = []

                    for i, row in enumerate(table.rows):
                        cells = [
                            ''.join(t.text for t in cell._element.iter() if t.tag.endswith('}t')).strip()
                            for cell in row.cells
                        ]
                        if i == 0:
                            headers = cells
                        else:
                            row_dict = {}
                            for h, v in zip(headers, cells):
                                row_dict[h] = v
                            rows_data.append(row_dict)

                    result[current_main_section]["3.2 Risk Yönetimi"] = rows_data
                # 3.3 Araştırma olanakları
                elif "3.3 Araştırma olanakları" not in result[current_main_section]:
                    headers = []
                    rows_data = []

                    for i, row in enumerate(table.rows):
                        cells = [
                            ''.join(t.text for t in cell._element.iter() if t.tag.endswith('}t')).strip()
                            for cell in row.cells
                        ]
                        if i == 0:
                            headers = cells
                        else:
                            row_dict = {}
                            for h, v in zip(headers, cells):
                                row_dict[h] = v
                            rows_data.append(row_dict)

                    result[current_main_section]["3.3 Araştırma olanakları"] = rows_data

            elif current_main_section == "4. YAYGIN ETKİ":
                if "Beklenen Etkiler" in result[current_main_section]:
                    continue  # sadece ilk tabloyu al

                headers = []
                rows_data = []

                for i, row in enumerate(table.rows):
                    cells = [
                        ''.join(t.text for t in cell._element.iter() if t.tag.endswith('}t')).strip()
                        for cell in row.cells
                    ]
                    if i == 0:
                        headers = cells
                    else:
                        row_dict = {}
                        for h, v in zip(headers, cells):
                            row_dict[h] = v
                        rows_data.append(row_dict)

                result[current_main_section]["Beklenen Etkiler"] = rows_data

            elif current_main_section == "BÜTÇE TALEP ÇİZELGESİ":
                if "Bütçe Kalemleri" in result[current_main_section]:
                    continue  # sadece ilk tabloyu al

                headers = []
                rows_data = []

                for i, row in enumerate(table.rows):
                    cells = [
                        ''.join(t.text for t in cell._element.iter() if t.tag.endswith('}t')).strip()
                        for cell in row.cells
                    ]
                    if i == 0:
                        headers = cells
                    else:
                        row_dict = {}
                        for h, v in zip(headers, cells):
                            row_dict[h] = v
                        rows_data.append(row_dict)

                result[current_main_section]["Bütçe Kalemleri"] = rows_data





    return result


# data = extract_full_2209a_structure("testWord.docx")
# import json
# print(json.dumps(data, indent=4, ensure_ascii=False))